"""python-docx extractor — for .docx files.

Reads paragraphs + table cell text. No styling, no images, no headers/
footers (for now — Sprint 22+ can extend). Plain text is enough for
search.
"""

from __future__ import annotations

from pathlib import Path

from hermes.memory.extractors import ExtractionResult, logger


def extract(path: Path) -> ExtractionResult:
    """Extract text from a .docx file via python-docx.

    Returns ExtractionResult with text=<paragraphs + table cells joined>,
    confidence=1.0, model='python-docx', error=None on success.

    Error cases:
    - python-docx not installed: error='python_docx_not_installed'
    - File not found: error='file_not_found'
    - Corrupted DOCX: error='parse_error: <reason>'
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        return ExtractionResult(
            text="",
            confidence=1.0,
            model="python-docx",
            error="python_docx_not_installed",
        )

    try:
        # python-docx.Document types as `str | IO[bytes] | None`; convert
        # Path to native str. `str(path)` on Windows gives backslashes
        # (which `open()` accepts), on Linux gives forward slashes. We
        # use `str(path)` not `as_posix()` because this path is consumed
        # by `python-docx` -> Python's `open()` (not persisted to DB/JSON),
        # so §7's POSIX convention does not apply here. as_posix() would
        # also work (open() handles `/` on Windows), but `str(path)` is
        # simpler and avoids any cross-platform concern.
        doc = Document(str(path))
    except FileNotFoundError:
        return ExtractionResult(
            text="", confidence=1.0, model="python-docx", error="file_not_found"
        )
    except Exception as exc:
        logger.warning(
            "python_docx_open_error",
            extra={"path": path.as_posix(), "error": str(exc)},
        )
        return ExtractionResult(
            text="",
            confidence=1.0,
            model="python-docx",
            error=f"parse_error: {exc}",
        )

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(parts)

    return ExtractionResult(text=text, confidence=1.0, model="python-docx", error=None)
